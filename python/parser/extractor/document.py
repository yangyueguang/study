# coding=utf-8
# author: Super
# email: 2829969299@qq.com
import fitz
import json
import abc
import os
import re
import sys
import cv2
import PIL
import math
import base64
import codecs
import hashlib
import numpy as np
import pandas as pd
from io import BytesIO
from skimage import morphology
import matplotlib.pyplot as plt
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches, Pt, Cm
from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
ocr_engine = None
try:
    from paddleocr import PaddleOCR
    ocr_engine = PaddleOCR(use_angle_cls=True, lang="ch", use_gpu=True)
except Exception as e:
    print('please install paddleocr')


class Dict(dict):
    def __init__(self, sdict=None):
        super(dict, self).__init__()
        sdict = sdict or {}
        for sk, sv in sdict.items():
            if isinstance(sv, dict):
                self[sk] = Dict(sv)
            else:
                self[sk] = sv

    def __getattr__(self, name):
        try:
            return self[name]
        except KeyError:
            # raise AttributeError(name)
            return None


class utils(object):
    @staticmethod
    def get_sections(shadow, gap, shadow_gap=10):
        if len(shadow) <= 0: return []
        satisfy = np.where(shadow > min(shadow_gap, max(2, shadow.max() / 10)))
        indexs = np.append(np.where(np.diff(satisfy) > gap)[1] + 1, values=[sys.maxsize], axis=0)
        sections = np.split(satisfy[0], indexs, axis=0)[:-1]
        return sections

    @staticmethod
    def get_middles(shadow, gap, shadow_gap=10):
        sections = utils.get_sections(shadow, gap, shadow_gap)
        middles = [(sections[ind][-1] + sections[ind + 1][0]) // 2 for ind in range(len(sections) - 1)]
        return middles

    @staticmethod
    def find_x_middles(binary: np.ndarray, col_gap=3):
        x_shadow = binary.sum(axis=0)
        return utils.get_middles(x_shadow, col_gap)

    @staticmethod
    def find_y_bottoms(binary: np.ndarray, row_gap=3):
        y_shadow = binary.sum(axis=1)
        sections = utils.get_sections(y_shadow, row_gap, 10)
        middles = [sections[ind + 1][0] for ind in range(len(sections) - 1)]
        return middles

    @staticmethod
    def find_y_middles(binary: np.ndarray, row_gap=3):
        y_shadow = binary.sum(axis=1)
        return utils.get_middles(y_shadow, row_gap)

    @staticmethod
    def find_middles(binary: np.ndarray, col_gap=3, row_gap=3):
        return utils.find_y_middles(binary, row_gap), utils.find_x_middles(binary, col_gap)

    @staticmethod
    def get_text_boxs(img: np.ndarray, is_split=False):
        boxs = []
        if ocr_engine is None:
            return boxs
        h, w, _ = img.shape

        def recognize_img(recog_img, offset=0):
            result = ocr_engine.ocr(recog_img)
            for l in result:
                txt = l[-1][0]
                if not txt.strip():
                    continue
                x0, y0 = l[0][0]
                x1, y1 = l[0][2]
                cw = (x1 - x0) / len(txt)
                chars = [Char(c, x0 + i * cw, y0 + offset, x0 + (i + 1) * cw, y1 + offset) for i, c in enumerate(txt)]
                boxs.append(Box(chars, x0, y0 + offset, x1, y1 + offset))

        if h <= 1000 or not is_split:
            recognize_img(img)
        else:
            middle = round(h / 2)
            top = img[0:middle, 0:w]
            bottom = img[middle - 5:h, 0:w]
            recognize_img(top)
            recognize_img(bottom, middle - 5)
        return boxs

    # 模版匹配从大图里找到这个小图对应的位置
    @staticmethod
    def find_img(imgPath, litle_img):
        target = cv2.imread(imgPath)
        th, tw = litle_img.shape[:2]
        res = cv2.matchTemplate(target, litle_img, cv2.TM_CCOEFF)
        res_sorted = sorted(res.max(axis=1), reverse=True)
        res_dif = [0] * 150
        for i in range(150):
            res_dif[i] = (res_sorted[i] - res_sorted[i + 1]) * 100. / res_sorted[i + 1]
        max_lastIdx = res_dif.index(sorted(res_dif, reverse=True)[0])
        idx = np.argwhere(res >= res_sorted[max_lastIdx])
        idx_set = set(np.unique(idx[:, 0]))
        for i in range(len(idx)):
            if idx[i, 0] in idx_set:
                idx_set.remove(idx[i, 0])
                tl = (idx[i, 1], idx[i, 0])
                br = (tl[0] + tw, tl[1] + th)
                cv2.rectangle(target, tl, br, (0, 0, 0), 1)
        cv2.imwrite(imgPath, target)


class Serializable(object):
    __metaclass__ = abc.ABCMeta

    @abc.abstractmethod
    def json(self) -> dict:
        pass

    @property
    @abc.abstractmethod
    def text(self) -> str:
        pass

    @property
    def type(self) -> str:
        return self.__class__.__name__.lower()

    @classmethod
    def load(cls, j: Dict):
        pass

    @staticmethod
    def md5(string: str) -> str:
        md = hashlib.md5()
        md.update(string.encode())
        return md.hexdigest()

    @staticmethod
    def load_json(json_file: str):
        with codecs.open(json_file, 'r', encoding='utf-8') as f:
            json_dict = json.load(f, encoding='utf-8')
        return json_dict

    def cat_binary(self, binary):
        plt.imshow(binary)
        plt.show()

    def __str__(self):
        return self.text


class Box(Serializable):
    def __init__(self, chars: list, *args):
        super(Box, self).__init__()
        if len(args) >= 4:
            self.x, self.y, self.r, self.b = map(float, args[:4])
        elif chars:
            self.x = math.floor(min([c.x for c in chars]))
            self.r = math.ceil(max([c.r for c in chars]))
            self.y = math.floor(min([c.y for c in chars]))
            self.b = math.ceil(max([c.b for c in chars]))
        else:
            self.x = self.y = self.b = self.r = 0
        self.color = 0
        self.chars = chars or []
        self.row_span = 1
        self.col_span = 1

    @property
    def width(self) -> float:
        return abs(self.r - self.x)

    @width.setter
    def width(self, w: float):
        self.r = self.x + w

    @property
    def height(self) -> float:
        return abs(self.b - self.y)

    @height.setter
    def height(self, h: float):
        self.b = self.y + h

    @property
    def center(self):
        return (self.x + self.r) / 2, (self.y + self.b) / 2

    @property
    def text(self) -> str:
        return ''.join([c.str for c in self.chars])

    @property
    def is_line(self):
        return isinstance(self, Table) and self.width < 8 or self.height < 8

    @property
    def rect(self) -> tuple:
        return math.floor(self.x), math.floor(self.y), math.ceil(self.r), math.ceil(self.b)

    def include_point(self, x: float, y: float):
        self.x = min(self.x, x)
        self.y = min(self.y, y)
        self.r = max(self.r, x)
        self.b = max(self.b, y)

    def include_box(self, box):
        self.include_point(box.x, box.y)
        self.include_point(box.r, box.b)
        return self

    def intersect(self, b):
        if not self.is_intersect(b):
            return Box([])
        return Box([], max(self.x, b.x), max(self.y, b.y), min(self.r, b.r), min(self.b, b.b))

    def is_intersect(self, b):
        return min(self.r, b.r) > max(self.x, b.x) and min(self.b, b.b) > max(self.y, b.y)

    def json(self) -> dict:
        return {
            'x': math.floor(self.x),
            'y': math.floor(self.y),
            'r': math.ceil(self.r),
            'b': math.ceil(self.b),
            'str': self.text,
            'type': self.type
        }

    @classmethod
    def load(cls, j: Dict):
        cs = [Char(i, j.x, j.y, j.x + 10, j.y + 10) for i in j.str]
        box = cls(cs, j.x, j.y, j.r, j.b)
        box.row_span = 1 if j.row_span is None else j.row_span
        box.col_span = 1 if j.col_span is None else j.col_span
        return box

    def __getitem__(self, i):
        return (self.x, self.y, self.r, self.b)[i]

    def __lt__(self, other):
        return self.y < other.y or (self.y == other.y and self.x < other.x)

    def __pos__(self):
        return self

    def __neg__(self):
        return Box(self.chars, -self.x, -self.y, -self.r, -self.b)

    def __bool__(self):
        return not self.is_line

    def __nonzero__(self):
        return not self.is_line

    def __eq__(self, rect):
        return bool(self.x == rect.x and self.y == rect.y and self.r == rect.r and self.b == rect.b)

    def __abs__(self):
        return (self.r - self.x) * (self.b - self.y)

    def __add__(self, b):
        return Box(self.chars + b.chars, min(self.x, b.x), min(self.y, b.y), max(self.r, b.r), max(self.b, b.b))

    def __sub__(self, b):
        return Box(self.chars, self.x - b.x, self.y - b.y, self.r - b.x, self.b - b.y)

    def __mul__(self, m):
        return Box(self.chars, self.x * m, self.y * m, self.r * m, self.b * m)

    def __truediv__(self, m):
        return Box(self.chars, self.x / m, self.y / m, self.r / m, self.b / m)

    __div__ = __truediv__

    def __contains__(self, x):
        if hasattr(x, "__float__"):
            return x in tuple(self)
        elif isinstance(x, Box):
            return self.x <= x.x <= x.r <= self.r and self.y <= x.y <= x.b <= self.b
        elif len(x) == 4:
            return self.x <= x[0] <= x[2] <= self.r and self.y <= x[1] <= x[-1] <= self.b
        else:
            return self.x <= x[0] <= self.r and self.y <= x[1] <= self.b

    def __or__(self, b):
        return self.include_box(b)

    def __and__(self, b):
        return self.intersect(b)

    def __hash__(self):
        return hash(tuple(self))

    @classmethod
    def load_span(cls, span: dict):
        chars = []
        color = span['color']
        for char in span['chars']:
            c = Char(char['c'], *char['bbox'])
            c.font = span['font']
            c.color = color
            chars.append(c)
        box = cls(chars, *span['bbox'])
        box.color = span['color']
        return box


class BaseElement(Box):
    def __init__(self, page, lines: list, *args):
        super(BaseElement, self).__init__([c for l in lines for c in l.chars], *args)
        self.lines = lines
        self.page = page
        self.parent = None
        self.children = []
        self._blocks = []

    def get_blocks(self):
        if not self._blocks and self.page is not None and self.page.own is not None:
            blocks = self.page.own.getText('rawdict', clip=(self.x, self.y, self.r, self.b))
            self._blocks = [b.get('lines', []) for b in blocks['blocks']]
        return self._blocks

    def layout(self):
        self.lines = []
        blocks = self.get_blocks()
        for block in blocks:
            for line in block:
                boxs = []
                for span in line['spans']:
                    boxs.append(Box.load_span(span))
                if not ''.join([b.text for b in boxs]).replace('\n', '').strip():
                    continue
                if self.lines and (line['bbox'][-1] + line['bbox'][1]) / 2 < self.lines[-1].b:
                    last_line = self.lines[-1]
                    last_line.include_box(Box(None, *line['bbox']))
                    last_line.boxs.extend(boxs)
                    last_line.fresh()
                    self.lines[-1] = last_line
                else:
                    lin = TextLine(boxs, *line['bbox'])
                    self.lines.append(lin)
        self.chars = [c for l in self.lines for c in l.chars]

    @property
    def global_y(self) -> float:
        return (self.page.y if self.page else 0) + self.y

    @classmethod
    def load(cls, j: Dict):
        rect = [j.x, j.y, j.r, j.b]
        text_line = TextLine([Box([Char(i, *rect) for i in j.str], *rect)], *rect)
        p = cls(None, [text_line], j.x, j.y, j.r, j.b)
        return p

    @property
    def text(self) -> str:
        if not self.chars:
            self.chars = [c for l in self.lines for c in l.chars]
        return ''.join([c.str for c in self.chars])


class Char(Box):
    __slots__ = ('x', 'y', 'r', 'b', 'font', 'color', 'str')

    def __init__(self, value: str = '', *args):
        super(Char, self).__init__(None, *args)
        self.color = 0
        self.font = None
        self.str = value

    @property
    def text(self) -> str:
        return self.str

    def __str__(self):
        return json.dumps(self.json(), ensure_ascii=False)

    def __eq__(self, other):
        return super(Char, self).__eq__(other) and self.str == other.str


class TextLine(Box):
    def __init__(self, boxs: [Box], *args):
        super(TextLine, self).__init__([c for box in boxs for c in box.chars], *args)
        self.boxs = boxs

    def fresh(self):
        self.chars = [c for i in self.boxs for c in i.chars]

    def __eq__(self, other):
        return self.text == other.text

    def __cmp__(self, other):
        if self.y == other.y:
            return self.x - other.x
        return self.y - other.y

    def __lt__(self, other):
        if self.y == other.y:
            return self.x < other.x
        return self.y < other.y

    def __hash__(self):
        return hash(self.text)

    def json(self) -> dict:
        boxs = [b.json() if b.row_span==b.col_span==1 else dict(b.json(), **{'row_span': b.row_span, 'col_span': b.col_span}) for b in self.boxs]
        res = super(TextLine, self).json()
        res.pop('str')
        res['boxs'] = boxs
        return res

    @classmethod
    def load(cls, j: Dict):
        return cls([Box.load(Dict(i)) for i in j.boxs], j.x, j.y, j.r, j.b)


class Table(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Table, self).__init__(page, lines, *args)
        self.rows = len(lines)
        self.cols = len(lines[0].boxs) if self.rows > 0 else 0
        self.page_number = 0
        self.grid = None
        self.binary = None
        self.boxs = []

    @property
    def cells(self) -> [Box]:
        return [cell for row in self.lines for cell in row.boxs]

    @property
    def matrix(self):
        matrix = [[cell.text for cell in line.boxs] for line in self.lines]
        res = pd.DataFrame(data=matrix)
        return res

    def add_edges(self):
        if self.grid is None:
            return
        self.grid[:3, :] = self.grid[-3:, :] = self.grid[:, :3] = self.grid[:, -3:] = 1

    def clip_edges(self):
        if self.grid is None:
            return
        if np.any(self.grid) != 1:
            self.add_edges()
            return
        rows = self.grid.sum(axis=1)
        cols = self.grid.sum(axis=0)
        leave_rows = np.where(rows > 0)[0]
        leave_cols = np.where(cols > 0)[0]
        self.r = self.x + leave_cols[-1]
        self.b = self.y + leave_rows[-1]
        self.x = self.x + leave_cols[0]
        self.y = self.y + leave_rows[0]
        self.grid = self.grid[leave_rows[0]: leave_rows[-1] + 1, leave_cols[0]:leave_cols[-1] + 1]
        self.grid[:3, :] = self.grid[-3:, :] = self.grid[:, :3] = self.grid[:, -3:] = 1

    def deal_sub_binary(self):
        grid_cells = self.find_cells_from_grid(self.grid)
        for row in grid_cells:
            row_middles_collection = []
            for c in row:
                l, t, r, b = c.rect
                cell_binary = self.binary[t:b, l:r]
                row_middles, col_middles = utils.find_middles(cell_binary, row_gap=1)
                for k in col_middles:
                    self.grid[t:b, l + k:l + k + 2] = 1
                if row_middles and row_middles[0] < 5:
                    row_middles.pop(0)
                row_middles_collection.append(row_middles)
                # for k in row_middles:
                #     self.binary[t+k:t+k+2, l:r] = 1
            # 上面横线不添加，从下面判断是否要添加
            row_lens = [len(rms) for rms in row_middles_collection]
            if self.page.is_ocr:
                row_middles_collection = [i for i in row_middles_collection if len(i) == max(row_lens)]
                row_lens = [len(rms) for rms in row_middles_collection]
            row_min_len = min(row_lens)
            if row_min_len > 0:
                l = row[0].rect[0]
                if row_min_len == max(row_lens):
                    row_middles_average = np.array(row_middles_collection).sum(axis=0) // len(row_lens)
                    for k in row_middles_average:
                        self.grid[t + k:t + k + 2, l:r] = 1
                elif row_min_len > 1:  # 如果各列行数不同，则以最小行数为基准，以各行顶部为基线
                    min_tops = []
                    for c in row:
                        cl, ct, cr, cb = c.rect
                        cell_binary = self.binary[ct:cb, cl:cr]
                        x_shadow = cell_binary.sum(axis=1)
                        sections = utils.get_sections(x_shadow, gap=1)
                        tops = [s[0] for s in sections][1:]
                        if not min_tops or len(min_tops) > len(tops):
                            min_tops = tops
                    for k in min_tops:
                        self.grid[t + k - 2:t + k, l:r] = 1

    def layout(self):
        boxs = self.boxs
        if not boxs:
            blocks = self.get_blocks()
            boxs = [Box.load_span(span) for block in blocks for line in block for span in line['spans']]
            boxs = [b for b in boxs if b.text.strip()]
        if self.is_line or self.grid is None or (not boxs and self.binary is None):
            self.boxs = []
            super().layout()
            return
        self.lines = []
        if self.binary is None:
            self.binary = np.zeros(self.grid.shape, dtype=int)
            for o in boxs:
                x, y, r, b = (o - self).rect
                offset = round((b - y) / 5)
                self.binary[y + offset:b - offset, x:r] = 1
        self.deal_sub_binary()
        grids = self.find_cells_from_grid(self.grid)
        self.rows = len(grids)
        self.cols = len(grids[0]) if grids else 0
        for r in grids:
            for c in r:
                c.x += self.x
                c.y += self.y
                c.r += self.x
                c.b += self.y
                for k in range(len(boxs) - 1, -1, -1):
                    span = boxs[k]
                    if span.center in c:
                        c.chars[0:0] = span.chars
                        boxs.pop(k)
            row = TextLine(r, self.x, r[0].y, self.r, r[0].b)
            self.lines.append(row)
        # deal empty row or col
        self.lines = [i for i in self.lines if i.chars]
        for j in range(self.cols - 1, -1, -1):
            col_chars = [c for l in self.lines for b in l.boxs[j:j + 1] for c in b.chars]
            if not col_chars:
                self.cols -= 1
                for l in self.lines:
                    l.boxs.pop(j)
        # deal row or col span
        for i, r in enumerate(self.lines):
            for j, b in enumerate(r.boxs):
                if b.row_span > 1 or b.col_span > 1:
                    for l in self.lines[i: i + b.row_span]:
                        be_merge_cells = l.boxs[j: j + b.col_span]
                        for c in be_merge_cells:
                            if not c.chars:
                                c.col_span = 0
                                c.row_span = 0

    def find_cells_from_grid(self, binary: np.ndarray):
        grids = []
        if not (binary.shape[0] and binary.shape[-1]):
            return grids
        shadow_y = binary.sum(axis=1)
        shadow_x = binary.sum(axis=0)
        most_y = np.argmax(np.bincount(np.array(shadow_y.tolist())))
        most_x = np.argmax(np.bincount(np.array(shadow_x.tolist())))
        shadow_y[shadow_y <= most_y+5] = 0
        shadow_x[shadow_x <= most_x+5] = 0
        rows = np.where(np.diff(shadow_y) > shadow_y.max() / 5)[0] + 1
        cols = np.where(np.diff(shadow_x) > shadow_x.max() / 6)[0] + 1
        rows = np.insert(rows, 0, values=0, axis=0)
        cols = np.insert(cols, 0, values=0, axis=0)
        split_rows = np.split(rows, np.where(np.diff(rows) > 5)[0] + 1, axis=0)
        rows = [s[0] for s in split_rows]
        split_cols = np.split(cols, np.where(np.diff(cols) > 10)[0] + 1, axis=0)
        cols = []
        for s in split_cols:
            # 这种情况是小分段不合理，集合起来
            middle = s[0]
            if len(s) > 1:
                start, end = s[0], s[-1]
                cp_binary = np.copy(binary[:, start:end + 3])
                middle = (start + end) // 2
                if start == 0:
                    middle = start+1
                elif end >= binary.shape[1] - 3:
                    middle = end
                cp_shadow = cp_binary.sum(axis=1)
                for i, a in enumerate(cp_shadow):
                    binary[i:i + 1, start:end + 3] = 1 if a >= cp_shadow.max() / 3 else 0
                    binary[i:i + 1, middle - 1:middle + 1] = 1 if a > 0 else 0
            cols.append(middle)
        w = 5  # 线宽
        for i in range(len(rows) - 1):
            row_cells = []
            for j in range(len(cols) - 1):
                m, n = i, j
                right_open, bottom_open = True, True
                while right_open or bottom_open:
                    top, bottom, left, right = rows[m], rows[m + 1], cols[n], cols[n + 1]
                    right_open = np.any(binary[min(top + w, bottom - 1): bottom, right-w: right + w]) != 1
                    bottom_open = np.any(binary[bottom: bottom + w, min(left + w, right - 1): right]) != 1
                    if bottom_open:
                        m += 1
                    if right_open:
                        n += 1
                rect = (cols[j], rows[i], cols[n + 1], rows[m + 1])
                cell = Box([], *rect)
                cell.row_span = m - i + 1
                cell.col_span = n - j + 1
                row_cells.append(cell)
            if row_cells:
                grids.append(row_cells)
        return grids

    @classmethod
    def load_from_metas(cls, page, meta_list: list, binary, x, y, r, b):
        t = cls(page, [], x, y, r, b)
        t.grid = binary
        t.add_edges()
        t.clip_edges()
        in_table_metas = []
        remove_indexs = []
        for inde in range(len(meta_list) - 1, -1, -1):
            m = meta_list[inde]
            if not isinstance(m, Graph) and m.center in t:
                in_table_metas.append(m)
                remove_indexs.append(inde)
        t.boxs = [b for p in in_table_metas for l in p.lines for b in l.boxs]
        t.layout()
        if len(t.cells) <= 2 and not t.is_line:
            t = Paragraph(t.page, t.lines, t.x, t.y, t.r, t.b)
        return t, remove_indexs

    def json(self) -> dict:
        res = super(Table, self).json()
        res.pop('str')
        res['rows'] = [r.json() for r in self.lines]
        return res

    @classmethod
    def load(cls, j: Dict):
        return cls(None, [TextLine.load(Dict(i)) for i in j.rows], j.x, j.y, j.r, j.b)

    @property
    def html(self) -> str:
        def get_row(c: Box):
            td = '<td' + (f' rowspan={c.row_span}' if c.row_span > 1 else '')
            return td + (f' colspan={c.col_span}' if c.col_span > 1 else '') + f'>{c.text}</td>\n'

        web = '<table>\n<tbody>'
        for row in self.lines:
            cells = [get_row(c) for c in row.boxs if c.row_span and c.col_span]
            web += f'<tr>\n{"".join(cells)}\n</tr>'
        web += '</tbody>\n</table>'
        return web


class Line(Table):
    def __init__(self, page, _, *args):
        super(Line, self).__init__(page, [], *args)


class Title(BaseElement):
    def __init__(self, page, lines: list, *args):
        super(Title, self).__init__(page, lines, *args)
        self.level = len(re.split(r'\.', re.search(r'([\d\.]*)', self.text).group(0).strip('.')))


class Paragraph(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Paragraph, self).__init__(page, lines, *args)

    @classmethod
    def load_paragraph(cls, page=None, lines: list = None, *args):
        p = cls(page, lines or [], *args)
        bold_or_color = [c for c in p.chars if c.color > 0 or 'Bold' in c.font]
        if bold_or_color and len(p.chars) < 40 and re.match(r'[\d](\.\d+)*\.?\s.{1,40}', p.text.strip()):
            return Title(page, lines, *args)
        return p


class Header(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Header, self).__init__(page, lines, *args)


class Footer(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Footer, self).__init__(page, lines, *args)


class Catelog(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Catelog, self).__init__(page, lines, *args)


class Graph(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Graph, self).__init__(page, lines, *args)
        self.img_id = None

    def save(self, file_name: str):
        if self.page is None:
            return
        ext = file_name.rsplit('.')[-1]
        ext = 'jpeg' if ext == 'jpg' else ext
        pix = self.get_pixmap()
        if pix.colorspace and pix.colorspace.n > 3 and ext == 'png':
            tmp_name = file_name + '.psd'
            try:
                pix.writeImage(tmp_name, 'psd')
                with PIL.Image.open(tmp_name) as pic:
                    pic.convert('RGBA').save(file_name, ext)
            finally:
                os.remove(tmp_name)
        else:
            pix.writeImage(file_name, ext)

    def show(self):
        if self.page is None:
            return
        pix = self.get_pixmap()
        pic = PIL.Image.open(BytesIO(pix.getPNGData()))
        pic.show(self.text)

    @property
    def is_real_graph(self):
        image_array = np.frombuffer(self.get_pixmap().getPNGData(), dtype=np.uint8)
        gray = cv2.imdecode(image_array, cv2.IMREAD_GRAYSCALE)
        binary = cv2.adaptiveThreshold(~gray, 1, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, -10)
        er = cv2.erode(binary, cv2.getStructuringElement(cv2.MORPH_RECT, (50, 1)), iterations=1)
        row = morphology.dilation(er, morphology.rectangle(3, 1))
        ec = cv2.erode(binary, cv2.getStructuringElement(cv2.MORPH_RECT, (1, 50)), iterations=1)
        ec = cv2.dilate(ec, cv2.getStructuringElement(cv2.MORPH_RECT, (1, 50)), iterations=1)
        col = morphology.dilation(ec, morphology.rectangle(1, 3))
        grid = row + col
        h, w = grid.shape
        cols = np.where(np.diff(np.where(grid.sum(axis=0) > h*0.8)) > 8)[1]
        rows = np.where(np.diff(np.where(grid.sum(axis=1) > w*0.8)) > 8)[1]
        return len(rows) <= 0 or len(cols) <= 0

    def get_pixmap(self):
        if self.page is None:
            return
        if self.img_id:
            # base_image = self.page.doc.extractImage(self.img_id)
            # image = PIL.Image.open(BytesIO(base_image["image"]))
            return fitz.Pixmap(self.page.doc, self.img_id)
        clip = [i/self.page.scale for i in self] if self.page.is_ocr else self.rect
        return self.page.getPixmap(clip=clip)


class Page(Box):
    def __init__(self, doc: fitz.Document, index: int, *args):
        super(Page, self).__init__([], *args)
        self.meta_list = []
        self.doc = doc
        self.index = index
        self.rotate = 0
        self.scale = 1
        self.own = None if self.doc is None else self.doc[self.index]
        self.is_ocr = False
        self.grid = None  # 网格线二进制
        self.binary = None  # 实体二进制

    def json(self) -> dict:
        res = super(Page, self).json()
        res.pop('str')
        res.update({
            'scale': self.scale,
            'index': self.index,
            'rotate': self.rotate,
            'meta_list': [m.json() for m in self.meta_list]
        })
        return res

    @classmethod
    def load(cls, j: Dict):
        p = cls(None, j.index, j.x, j.y, j.r, j.b)
        for m in j.meta_list:
            sub_cls = eval(m.get('type', 'paragraph').capitalize())
            obj = sub_cls.load(Dict(m))
            p.meta_list.append(obj)
        p.scale = j.scale
        p.rotate = j.rotate
        return p

    def parse_ocr(self):
        image_array = np.frombuffer(self.getPixmap().getPNGData(), dtype=np.uint8)
        img = cv2.imdecode(image_array, cv2.IMREAD_ANYCOLOR)
        boxs = utils.get_text_boxs(img, is_split=True)
        op = ImageLayout(None, img=img, page=self)
        op.layout_parse()
        op.fill_boxs(boxs)
        self.meta_list = op.meta_list
        self.grid = op.grid
        self.binary = op.block_binary

    def parse(self):
        meta_list = self.get_base_metas()
        if self.is_ocr:
            self.parse_ocr()
            return self
        self.binary = np.zeros(self.grid.shape, dtype=int)
        contents = [i for i in meta_list if not i.is_line]
        for i in contents:
            for l in i.lines:
                for b in l.boxs:
                    x, y, r, b = b.rect
                    self.binary[y:b + 1, x: r + 1] = 1
        columns = self._split_columns(meta_list)
        result = [m for ms, binary in columns for m in self.deal_no_line(ms, binary)]
        for _, grid in columns[:-1]:
            self.grid += grid
        # 处理几个段落合并在一起的情况
        new_result = []
        for m in result:
            if isinstance(m, Paragraph):
                paras = []
                new_lines = []
                for i in m.lines:
                    new_lines.append(i)
                    if i.rect[2] < m.r - 20 or i.text.endswith('。 '):
                        paras.append(Paragraph.load_paragraph(self, new_lines, m.x, m.y, m.r, m.b))
                        new_lines = []
                if new_lines:
                    paras.append(Paragraph.load_paragraph(self, new_lines, m.x, m.y, m.r, m.b))
                new_result.extend(paras)
            elif isinstance(m, Table) and m.is_line:
                new_result.append(Line(self, [], m.x, m.y, m.r, m.b))
            else:
                new_result.append(m)
        self.meta_list = new_result
        self.fresh()

    def get_base_metas(self):
        page = self.own
        meta_list = []
        binary = np.zeros(page.rect.irect[2:][::-1], dtype=int)
        page_raw = page.getTextPage().extractRAWDICT()
        if not page_raw['blocks']:
            self.is_ocr = True
            return meta_list
        images = page.getImageList()
        draws = page.getDrawings()
        draw_boxs = []
        graph_boxs = []
        paragraphs = []

        for ig in images:
            image_box = page.getImageBbox(ig[-2])
            graph = Graph(self, [], image_box.x0, image_box.y0, image_box.x1, image_box.y1)
            graph.img_id = ig[0]
            graph_boxs.append(graph)

        for d in draws:
            r = d['rect'].irect
            if r.width + r.height < 5 or (d['fill'] or [0])[0] > 0 and (
                    0.1 < r.width / r.height < 10 or min(r.width, r.height) > 10):
                continue
            if r.width > 20 and r.height > 20:  # 这里表面是表格，其实是复杂的绘图，直接用图来处理
                graph = Graph(self, [], r.x0 - r.width / 5, r.y0 - r.height / 4 - 10, r.x1 + r.width / 5,
                              r.y1 + r.height / 5)
                for g in graph_boxs:
                    if g.intersect(graph):
                        g.include_box(graph)
                        break
                else:
                    graph_boxs.append(graph)
                continue
            binary[r[1]:r[3] + 1, r[0]:r[2] + 1] = 1
        pr, pc = np.where(binary != 0)
        pr.sort()
        pc.sort()
        ploy_row = np.split(pr, np.where(np.diff(pr) > 1)[0] + 1)
        ploy_col = np.split(pc, np.where(np.diff(pc) > 1)[0] + 1)
        for r in ploy_row:
            for c in ploy_col:
                t = Table(self, [], c[0], r[0], c[-1], r[-1])
                t.grid = binary[r[0]:r[-1] + 1, c[0]:c[-1] + 1]
                if np.any(t.grid) != 1:
                    continue
                t.clip_edges()
                draw_boxs.append(t)
        for block in page_raw['blocks']:
            bb = block['bbox']
            block_lines = block['lines']
            for t in draw_boxs:
                if bb in t:
                    t._blocks.append(block_lines)
                    break
            else:
                for g in graph_boxs:
                    if bb in g:
                        g._blocks.append(block_lines)
                        break
                else:
                    p = Paragraph(self, [], *bb)
                    p._blocks.append(block_lines)
                    paragraphs.append(p)
        meta_list.extend(draw_boxs)
        meta_list.extend(graph_boxs)
        meta_list.extend(paragraphs)
        for i in meta_list:
            i.layout()
        meta_list = [m for m in meta_list if not (isinstance(m, Paragraph) and not m.chars)]
        meta_list.sort()
        self.grid = binary
        return meta_list

    def _split_columns(self, meta_list):
        column_binary = np.zeros(self.grid.shape, dtype=int)
        column_contents = [i for i in meta_list if not isinstance(i, Graph)]
        for i in column_contents:
            x, y, r, b = i.rect
            column_binary[y:b + 1, x: r + 1] = 1
        page_row_shadow = column_binary.sum(axis=0)
        page_satisfy = np.where(page_row_shadow > page_row_shadow.max() / 10)[0]
        page_sections = np.where(np.diff(page_satisfy) > 5)[0]
        page_lines = [(page_satisfy[s] + page_satisfy[s + 1]) // 2 for s in page_sections]
        new_list = []
        for l in page_lines:
            sub_binary = np.zeros(self.grid.shape, dtype=int)
            sub_binary[:, :l + 1] = self.grid[:, :l + 1]
            new_list.append([[i for i in meta_list if i.center[0] <= l], sub_binary])
            meta_list = [i for i in meta_list if i.center[0] > l]
            self.grid[:, :l + 1] = 0
        new_list.append([meta_list, self.grid])
        return new_list

    # 无线表格检测
    def deal_no_line(self, meta_list, binary):
        block_binary = np.zeros(self.grid.shape, dtype=int)
        contents = [i for i in meta_list if not i.is_line]
        for i in contents:
            for l in i.lines:
                for b in l.boxs:
                    x, y, r, b = b.rect
                    block_binary[y:b + 1, x: r + 1] = 1
        new_tables = []
        pre_cols = 0
        top = 0
        bottom = 0
        is_start = False
        w = self.grid.shape[-1]
        left, right = w, 0
        pre_table_indexs = []
        remove_indexs = []
        is_start = True
        top = bottom = meta_list[0].rect[1]
        for i in meta_list:
            if i.is_line:
                bottom = i.rect[1]
                if not is_start:
                    top = bottom
                    is_start = True
                continue
            if is_start:
                block_row = block_binary[top: i.rect[-1] - 2, 0:w]
                combine_middles = utils.find_x_middles(block_row)
                sections = len(combine_middles)
                c_row = block_binary[min(i.rect[1], bottom) + 2: i.rect[-1] - 2, 0:w]
                c_middles = utils.find_x_middles(c_row)
                c_sections = len(c_middles)
                is_table = isinstance(i, Table) and not i.is_line
                if is_table:
                    c_sections = i.cols - 1
                if pre_cols > 0:
                    if not is_table and (pre_cols == c_sections == sections or (
                            c_sections == sections and pre_cols < sections)):  # 如果两个分区一样
                        top = min(top, i.rect[1])
                        left = min(left, i.rect[0])
                        right = max(right, i.rect[2])
                        # binary[top: i.rect[-1]+1, left: left+3] = 1
                        # binary[top: i.rect[-1]+1, right: right+3] = 1
                        pre_table_indexs = combine_middles
                    else:  # 上个表格结束，这个表格开始
                        if pre_table_indexs:
                            for k in pre_table_indexs:
                                binary[top: bottom + 1, k - 1: k + 2] = 1
                            pre_table_indexs = []

                        t, temp = Table.load_from_metas(self, meta_list, binary[top: bottom, left: right + 3], left, top,
                                                        right + 3, bottom)
                        remove_indexs.extend(temp)
                        new_tables.append(t)
                        left, top, right, _ = i.rect
                elif c_sections > 0:  # 之前没有表格现在有表格
                    left, _, right, _ = i.rect
                    top = i.rect[1]
                else:  # 两行都没有表格
                    top = i.rect[1]
                pre_cols = c_sections
                bottom = i.rect[-1]
                if is_table:
                    top = i.rect[-1]
                    bottom = i.rect[-1]
                    pre_cols = 0
        if top != bottom:
            for middle in combine_middles:
                binary[top: bottom, middle: middle + 3] = 1
            t, temp = Table.load_from_metas(self, meta_list, binary[top: bottom + 3, left: right + 3], left, top, right + 3,
                                            bottom + 3)
            remove_indexs.extend(temp)
            new_tables.append(t)
            # meta_list.append(t)
        remove_indexs = list(set(remove_indexs))
        remove_indexs.sort(reverse=True)
        for i in remove_indexs:
            meta_list.pop(i)
        meta_list.extend(new_tables)
        meta_list.sort()
        return meta_list

    def show(self):
        PIL.Image.open(BytesIO(self.getPixmap().getPNGData())).show(str(self.index))

    def save(self, file_name: str):
        self.getPixmap().writePNG(file_name)

    @property
    def text(self) -> str:
        return '\n'.join([m.text for m in self.meta_list])

    def fresh(self):
        self.chars = [c for l in self.meta_list for c in l.chars]

    def __str__(self):
        return f'{self.x, self.y, self.r, self.b, self.index}'

    def __cmp__(self, other):
        return self.index - other.index

    def __eq__(self, other):
        return self.index == other.index

    def __lt__(self, other):
        return self.index < other.index

    def __hash__(self):
        return hash(id(self))

    def drawRect(self, rect, color=None, fill=None):
        img = self.own.newShape()
        img.drawRect(fitz.Rect(rect))
        img.finish(1, color, fill)
        img.commit(overlay=True)

    def getText(self, option="text", clip=None, flags=None):
        return self.own.getText(option, clip, flags)

    def getTextPage(self, clip=None, flags=0):
        return self.own.getTextPage(clip, flags)

    def getPixmap(self, matrix=fitz.Matrix(3, 3).preRotate(0), clip=None):
        return self.own.getPixmap(matrix=matrix, clip=clip)

    def annots(self, types=None):
        return self.own.annots(types)

    def widgets(self, types=None):
        return self.own.widgets(types)

    def links(self, kinds=None):
        return self.own.links(kinds)

    def getLinks(self):
        return self.own.getLinks()

    def updateLink(self, lnk):
        return self.own.updateLink(lnk)

    def insertLink(self, lnk, mark=True):
        self.own.insertLink(lnk, mark)

    def insertText(self, point, text, fontsize=11, color=None, fill=None):
        return self.own.insertText(point, text, fontsize, color=color, fill=fill)

    def insertTextbox(self, rect, buffer, fontsize=11, color=None, fill=None):
        return self.own.insertTextbox(rect, buffer, fontsize=fontsize, color=color, fill=fill)

    def insertImage(self, rect, filename=None, pixmap=None, stream=None):
        return self.own.insertImage(rect, filename, pixmap, stream)

    def newShape(self):
        return fitz.utils.Shape(self.own)

    def searchFor(self, text, quads=False, clip=None):
        return self.own.searchFor(text, quads=quads, clip=clip)

    def writeText(self, rect=None, writers=None, color=None):
        return self.own.writeText(rect, writers, color=color)


class Document(fitz.Document):
    def __init__(self, file, password: str = None, **kwargs):
        filename, stream = None, None
        if isinstance(file, str):
            filename = file
        else:
            stream = file
        super(Document, self).__init__(filename=filename, stream=stream, filetype='pdf')
        if self.needsPass:
            self.authenticate(password or '')
        self.metadata['name'] = self.name
        self.metadata['hash'] = hashlib.sha256(open(file, 'rb').read() if filename else stream).hexdigest() if file else None
        if not self.isPDF or self.needsPass and not password:
            raise ValueError('not pdf file or need password')
        global_y = 0
        self.pages = []
        for page_num, page in enumerate(self):
            p = Page(self, page.number, *page.rect.irect)
            p.b += global_y
            p.y += global_y
            p.rotate = page.rotation
            global_y += p.height
            self.pages.append(p)

    @classmethod
    def load_from_images(cls, imgs: list):
        doc = cls(None)
        for i in imgs:
            p = cls(i)
            rect = p[0].rect
            imgpdf = cls(p.convertToPDF())
            p.close()
            # doc.insertPDF(imgpdf)
            page = doc.newPage(width=rect.width, height=rect.height)
            page.showPDFpage(rect, imgpdf, 0)
        return doc

    def parse(self):
        for p in self.pages:
            print(p.index)
            p.parse()

    def save_layout(self, layout_path: str):
        for page in self:
            img = page.newShape()
            disp = fitz.Rect(page.CropBoxPosition, page.CropBoxPosition)
            img.drawRect(page.rect + disp)
            img.finish(color=None, fill=None)
            blks = page.getTextBlocks()
            for b in blks:
                img.drawRect(fitz.Rect(b[:4]) + disp)
                color = (1, 0 if b[-1] == 1 else 1, 0)
                img.finish(width=2, color=color)
            draws = page.getDrawings()
            for j in draws:
                img.drawRect(j['rect'].irect)
                img.finish(width=2, color=(1, 0, 0))
            image_list = page.getImageList()
            for ig in image_list:
                image_box = page.getImageBbox(ig[-2])
                img.drawRect(image_box.irect)
                img.finish(width=2, color=(0, 1, 0))
            img.commit()
        self.save(layout_path, garbage=4, deflate=True, clean=True)

    def json(self) -> dict:
        serialized_document = {
            'metadata': self.metadata,
            'pages': [page.json() for page in self.pages]
        }
        return serialized_document

    @property
    def text(self):
        return '\f'.join([p.text for p in self.pages])

    @classmethod
    def load(cls, json_dict):
        page_list = [Page.load(Dict(j)) for j in json_dict['pages']]
        document = cls(None)
        document.pages = page_list
        document.metadata = json_dict['metadata']
        return document

    def save(self, filename, garbage=0, deflate=0, clean=0):
        super(Document, self).save(filename, garbage=garbage, deflate=deflate, clean=clean)

    def getToC(self):
        return super(Document, self).getToC()

    def setToC(self, toc):
        super(Document, self).setToC(toc)

    def PDFCatalog(self):
        return super(Document, self).PDFCatalog()

    def newPage(self, pno=-1, width=595, height=842):
        return super(Document, self).newPage(pno, width, height)

    def insertPDF(self, doc: fitz.Document, from_page=-1):
        super(Document, self).insertPDF(doc, from_page)

    def remove_hidden(self):
        def remove_hidden(cont_lines):
            out_lines = []
            in_text = False
            suppress = False
            make_return = False
            for line in cont_lines:
                if line == b"BT":  # start of text object
                    in_text = True  # switch on
                    out_lines.append(line)  # output it
                    continue
                if line == b"ET":  # end of text object
                    in_text = False  # switch off
                    out_lines.append(line)  # output it
                    continue
                if line == b"3 Tr":  # text suppression operator
                    suppress = True  # switch on
                    make_return = True
                    continue
                if line[-2:] == b"Tr" and line[0] != b"3":
                    suppress = False  # text rendering changed
                    out_lines.append(line)
                    continue
                if line == b"Q":  # unstack command also switches off
                    suppress = False
                    out_lines.append(line)
                    continue
                if suppress and in_text:  # suppress hidden lines
                    continue
                out_lines.append(line)
            if make_return:
                return out_lines
            else:
                return None

        for p in self:
            xref = p.getContents()[0]  # only one b/o cleaning!
            cont = self.xrefStream(xref)
            cont_lines = remove_hidden(cont.splitlines())  # remove hidden text
            if cont_lines:
                cont = b"\n".join(cont_lines)
                self.updateStream(xref, cont)

    # 写入word文本文档
    def save_to_docx(self, name):
        doc = DocxDocument()
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        section = doc.sections[0]
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        def add_paragraph(p: Paragraph, color: RGBColor, font_size: int = 12):
            paragraph = doc.add_paragraph(p.text)
            st = paragraph.add_run()
            st.bold = False
            st.italic = False
            st.underline = False
            st.font.size = Pt(font_size)
            st.font.color.rgb = color
            fm = paragraph.paragraph_format
            fm.left_indent = Inches(0)
            fm.right_indent = Inches(0)
            fm.first_line_indent = Pt(2 * font_size)
            fm.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        def add_table(t: Table):
            docx_table = doc.add_table(rows=t.rows, cols=t.cols, style='Table Grid')
            docx_table.alignment = WD_TAB_ALIGNMENT.CENTER
            for i, row in enumerate(t.lines):
                for j, b in enumerate(row.boxs):
                    docx_table.rows[i].cells[j].text = b.text
                    docx_table.rows[i].cells[j].width = Pt(b.width)

        def rgb(color) -> RGBColor:
            a = color//256//256
            b = color%(256*256)//256
            c = color%256
            return RGBColor(min(255, a), min(255, b), min(255, c))

        for p in self.pages:
            for m in p.meta_list:
                if isinstance(m, Paragraph) and m.chars:
                    font_size = m.chars[0].height
                    color = m.chars[0].color
                    add_paragraph(m, rgb(color), font_size)
                elif isinstance(m, Catelog):
                    for l in m.lines:
                        add_paragraph(l, RGBColor(255, 255, 255), l.chars[0].height)
                elif isinstance(m, Title):
                    st = doc.add_heading(m.text, m.level).add_run()
                    st.bold = True
                    st.font.color.rgb = rgb(m.chars[0].color)
                elif isinstance(m, Table):
                    add_table(m)
                elif isinstance(m, Graph) and m.page is not None:
                    pix = m.get_pixmap()
                    img = BytesIO(pix.getPNGData())
                    doc.add_picture(img, width=Pt(m.width), height=Pt(m.height))
                    if m.text.strip():
                        add_paragraph(m, RGBColor(255, 255, 255), m.chars[0].height)
                else:
                    print('')
            doc.add_page_break()
        doc.save(name)

    def html(self) -> str:
        def template() -> str:
            tmp = """
              <!DOCTYPE html><head><meta charset="UTF-8"><title></title><style>
                  h1{text-align: center}
                  p {text-indent:2em;}
                  th {background-color: #5182BB;text-align: center;font-size: 14px;font-weight: bold;}
                  table, td {color: #000;}
                  tr, th {border:1px solid #5182BB;}
                  tr:nth-child(even) {background-color: #f2f2f2;}
                  tr:hover {background-color:#f5f5f5;}
                  td{padding: 5px 20px;font-size: 12px;font-family:Verdana; text-align: left;border:1px solid #5182BB;}
	              td[colspan]{text-align: center;}
                  table tbody {display: block;width: 1000px;overflow-x: scroll;border-collapse:collapse;}
              </style></head><body></body></html>
              """
            return tmp

        pages = []
        for idx, page in enumerate(self.pages):
            lines = [f'<h1>第{str(idx + 1)}页</h1>']
            for m in page.meta_list:
                w = ''
                if isinstance(m, Table):
                    w = m.html
                elif isinstance(m, Title):
                    w = f'<h{m.level+1}>{m.text}</h{m.level+1}>'
                elif isinstance(m, Paragraph):
                    w = f'<p>{m.text}</p>'
                elif isinstance(m, Graph) and m.page is not None:
                    pix = m.get_pixmap()
                    md = pix.getPNGData()
                    pic = base64.b64encode(md).decode('utf-8')
                    w = f'<img src="data:image/jpeg;base64, {pic}"/>'
                    if m.text.strip():
                        w += f'<p>{m.text}</p>'
                elif isinstance(m, Line):
                    w = f'<div style="border:1px solid #CCC"></div>'
                lines.append(w)
            pages.append('\n'.join(lines))
        html_body = f'\n<div style="border:2px solid #FF0000"></div>\n'.join(pages)
        html = template().replace(r'<body></body>', html_body.replace('\\', ''))
        return html


class ImageLayout(Serializable):
    def __init__(self, img_path, img=None, page=None):
        self.img = img if img is not None else cv2.imread(img_path, 1)
        gray = cv2.cvtColor(self.img, cv2.COLOR_BGR2GRAY)
        gray[:10, :] = gray[-10:, :] = gray[:, :10] = gray[:, -10:] = 255  # 删除边界10像素
        self.binary = cv2.adaptiveThreshold(~gray, 1, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, -10)
        self.rows, self.cols = self.binary.shape
        self.meta_list = []
        self.page = page
        if self.page is not None:
            self.page.scale = self.rows / page.height
        self.grid = None
        self.block_binary = None

    def find_graph(self, gray: np.ndarray) -> list:
        height, width = gray.shape
        gray = gray.copy()
        line = gray.copy()
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (50, 1))
        line = cv2.erode(line, kernel, iterations=1)
        line = cv2.dilate(line, kernel, iterations=1)
        line = morphology.dilation(line, morphology.rectangle(35, 1))
        line = morphology.erosion(line, morphology.rectangle(1, 10))
        line = morphology.dilation(line, morphology.rectangle(1, 10))
        gray = morphology.dilation(gray, morphology.rectangle(1, 50))
        gray = morphology.erosion(gray, morphology.rectangle(10, 1))
        gray = morphology.dilation(gray, morphology.rectangle(10, 1))
        contours, hierarchy = cv2.findContours(gray + line, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
        # cv2.drawContours(img, contours, -1, (255, 0, 0), 2)
        areas = []
        for c in contours[::-1]:
            # area = cv2.contourArea(c)
            # box = np.int0(cv2.boxPoints(cv2.minAreaRect(c)))
            # cv2.fillConvexPoly(img, c, (255, 0, 0))
            # cv2.drawContours(img, [box], 0, (0, 0, 255), 2)
            x, y, w, h = cv2.boundingRect(c)
            if h > height / 15 and w > width / 10:
                sub_img = gray[y:y+h, x:x+w]
                cct, cch = cv2.findContours(sub_img, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_NONE)
                for cc in cct:
                    cx, cy, cw, ch = cv2.boundingRect(cc)
                    if ch > height / 15:
                        areas.append([x, y, x + w, y + h])
                        break
            elif width / 10 < w < width / 1.5:
                for a in areas:
                    if 0 < y - a[-1] < h and a[0] <= x <= x+w <= a[2]:  # 应该在图下面的标题
                        a[-1] = y+h
                        break
        return areas

    def split_columns(self, binary):
        block_binary = morphology.dilation(binary, morphology.rectangle(30, 1))
        block_binary = morphology.dilation(block_binary, morphology.rectangle(1, 10))
        page_row_shadow = block_binary.sum(axis=0)
        page_lines = [0] + utils.get_middles(page_row_shadow, 10, block_binary.shape[0]/16) + [block_binary.shape[-1]]
        for i in range(len(page_lines)-1, 0, -1):
            if page_lines[i] - page_lines[i-1] < block_binary.shape[-1]/5:
                page_lines.pop(i-1)
        sections = [[page_lines[i], page_lines[i+1]] for i in range(len(page_lines)-1)]
        return [self] if len(sections) <= 1 else [ImageLayout(None, self.img[:, l:r], self.page) for l, r in sections]

    # 识别横线
    def recognize_row(self):
        scale, scale1 = 20, 300
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (self.cols // scale, 1))
        eroded = cv2.erode(self.binary, kernel, iterations=1)
        img_row = cv2.dilate(eroded, kernel, iterations=1)
        col_selem = morphology.rectangle(self.cols // scale1, 1)
        self.grid_row = morphology.dilation(img_row, col_selem)

    # 识别竖线
    def recognize_column(self):
        scale, scale1 = 20, 300
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, self.rows // scale))
        eroded = cv2.erode(self.binary, kernel, iterations=1)
        img_col = cv2.dilate(eroded, kernel, iterations=1)
        row_selem = morphology.rectangle(1, self.rows // scale1)
        self.grid_col = morphology.dilation(img_col, row_selem)

    # 顶点检测
    def recognize_point(self):
        point_img = cv2.bitwise_and(self.grid_row, self.grid_col)
        ys, xs = np.where(point_img > 0)
        myxs = self._consecutive(xs)
        myys = self._consecutive(ys)
        points = [((x[0] + x[-1])//2, (y[0] + y[-1])//2) for x in myxs for y in myys]
        return points

    # 标识表格
    def recognize_table(self):
        self.grid = cv2.bitwise_or(self.grid_row, self.grid_col)
        block_img = cv2.subtract(self.binary, self.grid)
        # 降噪之后横向加粗纵向锐化再纵向加粗形成文字block
        block_img = cv2.GaussianBlur(block_img, (5, 5), 0)
        block_img = morphology.dilation(block_img, morphology.rectangle(1, 25))
        block_img = morphology.erosion(block_img, morphology.rectangle(5, 1))
        self.block_binary = morphology.dilation(block_img, morphology.rectangle(10, 1))
        self.grid = self.grid.astype(np.int)

    # 对数组进行非连续分组
    def _consecutive(self, data, stepsize=1):
        data.sort()
        return np.split(data, np.where(np.diff(data) > stepsize)[0] + 1)

    # 有线框和外边框
    def find_tables(self):
        draw_boxs = []
        pr, pc = np.where(self.grid != 0)
        pr.sort()
        pc.sort()
        ploy_row = np.split(pr, np.where(np.diff(pr) > 1)[0] + 1)
        ploy_col = np.split(pc, np.where(np.diff(pc) > 1)[0] + 1)
        last_t = None
        for r in ploy_row:
            for c in ploy_col:
                if not len(r) or not len(c) or np.bincount(r).max(initial=None) < self.cols/4:
                    continue
                t = Table(self.page, [], c[0], r[0], c[-1], r[-1])
                t.grid = self.grid[r[0]:r[-1] + 1, c[0]:c[-1] + 1]
                t.binary = self.block_binary[r[0]:r[-1]+1, c[0]:c[-1]+1]
                if np.any(t.grid) != 1:
                    continue
                if last_t is not None and t.is_line:
                    top, bottom = last_t.rect[1], t.rect[-1]
                    left, right = min(last_t.rect[0], t.rect[0]), max(last_t.rect[2], t.rect[2])
                    current_row = self.block_binary[last_t.rect[-1]:bottom, left:right]
                    lines = [0] + utils.find_y_middles(current_row) + [bottom]
                    sections = [[lines[i], lines[i+1]] for i in range(len(lines)-1)]
                    section_middles = [len(utils.find_x_middles(current_row[s[0]:s[1], :])) for s in sections]
                    c_middle = utils.find_x_middles(current_row)
                    com_grid = self.grid[top:bottom + 1, left:right + 1]
                    com_binary = self.block_binary[top:bottom + 1, left:right + 1]
                    middle = utils.find_x_middles(com_binary)
                    if c_middle:  # 说明上个线跟这个线之间有表格
                        total = len(section_middles)
                        good_mid = len([i for i in section_middles if i >= len(c_middle)])
                        if good_mid < total*2/3 or section_middles.count(0) > total/3:
                            draw_boxs.append(t)
                            last_t = t
                            continue
                        if middle:  # 说明上个表格跟这个线之间有表格而且两个表格分区一致 先清除竖线后加不至于过粗
                            cg_shadow = com_grid.sum(axis=1)
                            for i, a in enumerate(cg_shadow):
                                com_grid[i:i + 1, :] = 1 if a >= cg_shadow.max() / 3 else 0
                            for k in middle:
                                com_grid[:, k:k + 2] = 1
                        else:
                            for k in c_middle:
                                com_grid[last_t.rect[-1]:bottom, k:k+2] = 1
                        last_t.grid = com_grid
                        last_t.binary = com_binary
                        last_t.x, last_t.y, last_t.r, last_t.b = left, top, right, bottom
                        last_t.add_edges()
                        last_t.layout()
                        continue
                t.add_edges()
                t.layout()
                draw_boxs.append(t)
                last_t = t
        return draw_boxs

    # 处理无线表格
    def layout_no_line(self, block_binary):
        block_binary = morphology.erosion(block_binary, morphology.rectangle(5, 1))
        block_binary = morphology.dilation(block_binary, morphology.rectangle(1, 15))
        page_sections = utils.get_sections(block_binary.sum(axis=0), 10)
        left, right = page_sections[0].min(), page_sections[-1].max()
        # row_middles = utils.find_y_middles(block_binary, row_gap=5) + [self.rows]
        row_middles = utils.find_y_bottoms(block_binary, row_gap=5) + [self.rows]
        top, bottom = 0, 0
        pre_middles = []
        meta_list = []
        row_nums = 0
        text_liness = []

        def new_table(top, bottom, rl, rr, grid=self.grid, page=self.page):
            height, width = grid.shape
            if bottom - top < max(height / 50, width / 20) or row_nums <= 1:
                new_text_line(left, top, right, bottom)
                return
            t = Table(page, [], 0, top, width, bottom)
            t.binary = block_binary[top:bottom + 1, 0:width]
            for m in pre_middles:
                grid[top:bottom + 1, m:m + 1] = 1
            t.grid = grid[top:bottom + 1, 0:width]
            t.add_edges()
            meta_list.append(t)

        def new_text_line(x, y, r, b, page=self.page):
            gap = max(20, (right - left) / 20)
            line = TextLine([], x, y, r, b)
            if x - left > gap and text_liness:
                p = Paragraph(page, [i for i in text_liness], left, text_liness[0].y, right, y)
                meta_list.append(p)
                text_liness.clear()
            text_liness.append(line)
            if right - r > gap or line.height > row_middles[-1] / 15:
                p = Paragraph(page, [i for i in text_liness], left, text_liness[0].y, right, b)
                meta_list.append(p)
                text_liness.clear()

        for i, k in enumerate(row_middles):
            row_nums += 1
            row = block_binary[bottom: k, 0: self.cols]
            row_sections = utils.get_sections(row.sum(axis=0), 3)
            col_middles = [(row_sections[ind][-1] + row_sections[ind + 1][0]) // 2 for ind in range(len(row_sections) - 1)]
            if len(row_sections[0]) <= 0:
                row_nums -= 1
                continue
            rl, rr = row_sections[0].min(), row_sections[-1].max()
            # col_middles = utils.find_x_middles(row, 3)
            combine = block_binary[top: k, 0: self.cols]
            com_middles = utils.find_x_middles(combine, 3)
            merge_pre = 0 < len(pre_middles) <= len(col_middles) == len(com_middles)
            merge_col = 0 < len(col_middles) <= len(pre_middles) == len(com_middles)
            if len(col_middles) > 0 and (merge_pre or merge_col):
                if row_nums == 1:
                    self.grid[bottom:bottom + 1, left:right] = 1
                bottom = k
                pre_middles = com_middles
                self.grid[k:k+1, left:right] = 1
                continue
            if len(pre_middles) > 0:
                row_nums -= 1
                new_table(top, bottom, rl, rr)
                row_nums = 0
            top = bottom
            if len(col_middles) > 0:
                bottom = k
            else:
                new_text_line(rl, bottom, rr, k)
                top = bottom = k
                row_nums = 0
            pre_middles = col_middles
        if top < bottom:
            new_table(top, bottom, left, right/2)
        if text_liness:
            p = Paragraph(self.page, text_liness, left, text_liness[0].y, right, self.rows)
            meta_list.append(p)
        return meta_list

    # 布局分析
    def layout_parse(self):
        meta_list = []
        graphs = self.find_graph(self.binary)
        meta_graphs = []
        for g in graphs:
            graph = Graph(self.page, [], *g)
            if graph.is_real_graph:
                meta_graphs.append(graph)
                self.img[g[1]:g[-1]+1, g[0]:g[2]+1] = 255
                self.binary[g[1]:g[-1]+1, g[0]:g[2]+1] = 0
        meta_list.extend(meta_graphs)
        pages = self.split_columns(self.binary)
        offset = 0
        for p in pages:
            page_metas = []
            p.recognize_row()
            p.recognize_column()
            p.recognize_table()
            tables = p.find_tables()
            block_binary = p.block_binary.copy()
            for i in tables:
                i.x += offset
                i.r += offset
                page_metas.append(i)
                if i.is_line:
                    continue
                block_binary[i.rect[1]:i.rect[-1] + 1, i.rect[0]:i.rect[2] + 1 + 5] = 0
            metas = p.layout_no_line(block_binary)
            for i in metas:
                i.x += offset
                i.r += offset
                for tl in i.lines:
                    tl.x += offset
                    tl.r += offset
                for j in page_metas+meta_graphs:
                    if j.is_line:
                        continue
                    if j.y < i.b < j.b:
                        i.b = j.y - 1
                    elif j.y <= i.y < j.b:
                        i.y = j.b + 1
                    elif j.center in i:
                        i.b = j.y - 1
            page_metas.extend(metas)
            meta_list.extend(page_metas)
            if offset == 0:
                meta_list.sort()
            offset += p.cols
        self.meta_list = meta_list

    def fill_boxs(self, boxs: list):
        for b in boxs:
            for m in self.meta_list:
                if b.center in m:
                    if isinstance(m, Paragraph):
                        for l in m.lines:
                            if b.center in l:
                                l.boxs.append(b)
                                break
                        else:
                            l.boxs.append(b)
                    elif isinstance(m, Table):
                        m.boxs.append(b)
                    elif isinstance(m, Graph):
                        line = TextLine([b], *b.rect)
                        m.lines.append(line)
                    break
            else:
                print('new paragraph')
        # 标题检测
        for i, m in enumerate(self.meta_list):
            if isinstance(m, Table):
                m.layout()
                if m.is_line:
                    self.meta_list[i] = Line(m.page, [], m.x, m.y, m.r, m.b)
            elif isinstance(m, Paragraph):
                for l in m.lines:
                    l.fresh()
                if len(m.lines) == 1 and re.match(r'\d+\s*((\.\s*\d+\s*)*\.?\s|(\.\d+)+[\.\s]*).{1,40}', m.text.strip()):
                    self.meta_list[i] = Title(m.page, m.lines, m.x, m.y, m.r, m.b)
        # 验证是否有没识别到的单元格
        empty_cells = [c for m in self.meta_list if isinstance(m, Table) for c in m.cells if not c.chars and c.col_span and c.row_span]
        for c in empty_cells:
            x, y, r, b = c.rect
            cell_img = self.img[y: b, x: r]
            sub_binary = cv2.resize(cell_img, (1000, int(1000/c.width*c.height)))
            cell_box = utils.get_text_boxs(sub_binary)
            c.chars = [char for b in cell_box for char in b.chars]
        empty_lines = [l for m in self.meta_list if isinstance(m, Paragraph) for l in m.lines if not l.boxs]
        for l in empty_lines:
            x, y, r, b = l.rect
            cell_img = self.img[y-b+y: b+b-y, x: r]
            if len(cell_img) <= 0:
                continue
            scale_height = int(1000 / l.width * l.height)
            cell_box = utils.get_text_boxs(cv2.resize(cell_img, (1000, scale_height*3)))
            chars = [char for b in cell_box for char in b.chars if scale_height < char.center[-1] < scale_height * 2]
            l.boxs.append(Box(sorted(chars, key=lambda c: c.x), *l.rect))
            l.fresh()
