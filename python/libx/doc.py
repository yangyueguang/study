# coding=utf-8
import re
import math
import fitz
import numpy as np
import pandas as pd
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.table import _Row, _Cell, Table
from docx.oxml import table, text, ns, OxmlElement
from docx.text.paragraph import Paragraph
from .document import Document


class Dict(dict):
    def __init__(self, sdict=None):
        super(dict, self).__init__()
        sdict = sdict or {}
        for sk, sv in sdict.items():
            if isinstance(sv, dict):
                self[sk] = Dict(sv)
            else:
                self[sk] = sv

    def __getattr__(self, name):
        try:
            return self[name]
        except KeyError:
            return None


class Serializable(object):

    def json(self) -> dict:
        return {}

    @property
    def text(self) -> str:
        return ''

    @property
    def type(self) -> str:
        return self.__class__.__name__.lower()

    @classmethod
    def load(cls, j: Dict):
        pass

    def cat_binary(self, binary: np.ndarray):
        """
        用于界面可视化查看grid或者binary
        """
        pass


class Box(Serializable):
    def __init__(self, chars: list, *args):
        super(Box, self).__init__()
        self.x, self.y, self.r, self.b = map(float, args[:4])
        self.chars = chars or []
        self.row_span = 1
        self.col_span = 1

    @property
    def width(self) -> float:
        return abs(self.r - self.x)

    @width.setter
    def width(self, w: float):
        self.r = self.x + w

    @property
    def height(self) -> float:
        return abs(self.b - self.y)

    @height.setter
    def height(self, h: float):
        self.b = self.y + h

    @property
    def center(self):
        return (self.x + self.r) / 2, (self.y + self.b) / 2

    @property
    def text(self) -> str:
        return ''.join([c.str for c in self.chars])

    @property
    def is_line(self):
        return isinstance(self, Table) and self.width < 8 or self.height < 8

    @property
    def rect(self) -> tuple:
        return math.floor(self.x), math.floor(self.y), math.ceil(self.r), math.ceil(self.b)

    def include_point(self, x: float, y: float):
        pass

    def include_box(self, box):
        pass

    def intersect(self, b):
        """取交集区域"""
        pass

    def is_intersect(self, b) -> bool:
        """是否有交集"""
        return False


class BaseElement(Box):
    def __init__(self, page, lines: list, *args):
        super(BaseElement, self).__init__([], *args)
        self.lines = lines
        self.page = page
        self.parent = None
        self.children = []

    @property
    def global_y(self) -> float:
        return (self.page.y if self.page else 0) + self.y


class Char(Box):
    def __init__(self, value: str = '', *args):
        super(Char, self).__init__([], *args)
        self.color = 0
        self.font = None
        self.str = value


class TextLine(Box):
    def __init__(self, boxs: [Box], *args):
        super(TextLine, self).__init__([], *args)
        self.boxs = boxs


class Table(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Table, self).__init__(page, lines, *args)
        self.rows = len(lines)
        self.cols = len(lines[0].boxs) if self.rows > 0 else 0
        self.page_number = 0
        self.grid: np.ndarray = None
        self.binary: np.ndarray = None
        self.boxs: [Box] = []

    @property
    def cells(self) -> [Box]:
        return [cell for row in self.lines for cell in row.boxs]

    @property
    def matrix(self) -> pd.DataFrame:
        matrix = [[cell.text for cell in line.boxs] for line in self.lines]
        res = pd.DataFrame(data=matrix)
        return res

    @property
    def html(self) -> str:
        return ''


class Line(Table):
    def __init__(self, page, _, *args):
        super(Line, self).__init__(page, [], *args)


class Title(BaseElement):
    def __init__(self, page, lines: list, *args):
        super(Title, self).__init__(page, lines, *args)
        self.level = len(re.split(r'\.', re.search(r'([\d\.]*)', self.text).group(0).strip('.')))


class Paragraph(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Paragraph, self).__init__(page, lines, *args)


class Header(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Header, self).__init__(page, lines, *args)


class Footer(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Footer, self).__init__(page, lines, *args)


class Catelog(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Catelog, self).__init__(page, lines, *args)


class Graph(BaseElement):
    def __init__(self, page, lines: [TextLine], *args):
        super(Graph, self).__init__(page, lines, *args)
        self.img_id = None

    def save(self, file_name: str):
        pass

    def show(self):
        pass


class Page(Box):
    def __init__(self, doc: fitz.Document, index: int, *args):
        super(Page, self).__init__([], *args)
        self.meta_list = []
        self.doc = doc
        self.index = index
        self.rotate = 0
        self.scale = 1
        self.own = self.doc[self.index]
        self.is_ocr = False
        self.grid = None  # 网格线二进制
        self.binary = None  # 实体二进制

    def parse(self):
        pass

    def show(self):
        pass

    def save(self, file_name: str):
        pass

    def drawRect(self, rect, color=None, fill=None):
        pass

    def getText(self, option="text", clip=None, flags=None):
        return self.own.getText(option, clip, flags)

    def getTextPage(self, clip=None, flags=0):
        return self.own.getTextPage(clip, flags)

    def getPixmap(self, matrix=fitz.Matrix(3, 3).preRotate(0), clip=None):
        return self.own.getPixmap(matrix, clip=clip)

    def annots(self, types=None):
        return self.own.annots(types)

    def widgets(self, types=None):
        return self.own.widgets(types)

    def links(self, kinds=None):
        return self.own.links(kinds)

    def getLinks(self):
        return self.own.getLinks()

    def updateLink(self, lnk):
        return self.own.updateLink(lnk)

    def insertLink(self, lnk, mark=True):
        self.own.insertLink(lnk, mark)

    def insertText(self, point, text, fontsize=11, color=None, fill=None):
        return self.own.insertText(point, text, fontsize, color=color, fill=fill)

    def insertTextbox(self, rect, buffer, fontsize=11, color=None, fill=None):
        return self.own.insertTextbox(rect, buffer, fontsize=fontsize, color=color, fill=fill)

    def insertImage(self, rect, filename=None, pixmap=None, stream=None):
        return self.own.insertImage(rect, filename, pixmap, stream)

    def newShape(self):
        return fitz.utils.Shape(self.own)

    def searchFor(self, text, quads=False, clip=None):
        return self.own.searchFor(text, quads=quads, clip=clip)

    def writeText(self, rect=None, writers=None, color=None):
        return self.own.writeText(rect, writers, color=color)


class ImageLayout(Serializable):
    def __init__(self, img_path, img=None, page: Page=None):
        pass

    def find_graph(self, gray: np.ndarray) -> list:
        return []

    # 布局分析
    def layout_parse(self):
        pass

    def fill_boxs(self, boxs: list):
        pass


class Pdf(Document):
    def __init__(self, file, password: str = None, **kwargs):
        self.pages: [Page] = []
        super(Doc, self).__init__(file, password, **kwargs)

    @classmethod
    def load_from_images(cls, imgs: list):
        super(Doc, cls).load_from_images(imgs)

    def parse(self):
        """pass every page"""
        super(Doc, self).parse()

    def save_layout(self, layout_path: str):
        super(Doc, self).save_layout(layout_path)

    def json(self) -> dict:
        return super(Doc, self).json()

    @classmethod
    def load(cls, json_dict):
        return super(Doc, cls).load(json_dict)

    def save(self, filename, garbage=0, deflate=0, clean=0):
        super(Doc, self).save(filename, garbage=garbage, deflate=deflate, clean=clean)

    def getToC(self):
        return super(Doc, self).getToC()

    def setToC(self, toc):
        super(Doc, self).setToC(toc)

    def PDFCatalog(self):
        return super(Doc, self).PDFCatalog()

    def newPage(self, pno=-1, width=595, height=842):
        return super(Doc, self).newPage(pno, width, height)

    def insertPDF(self, doc: fitz.Document, from_page=-1):
        super(Doc, self).insertPDF(doc, from_page)

    def remove_hidden(self):
        super(Doc, self).remove_hidden()

    # 写入word文本文档
    def save_to_docx(self, name):
        super(Doc, self).save_to_docx(name)

    def html(self) -> str:
        return super(Doc, self).html()


class Doc(docx.document.Document):
    def __init__(self, path):
        doc = docx.Document(path)
        super().__init__(doc._element, doc._part)
        self.path = path
        
    def items(self):
        for i in self.element.body.iterchildren():
            if isinstance(i, table.CT_Tbl):
                yield Table(i, self)
            elif isinstance(i, text.paragraph.CT_P):
                yield Paragraph(i, self)
        
    def add_table(self, rows, cols, style='Table_Grid', df=None):
        if df is None:
            return super(Doc, self).add_table(rows, cols, style)
        table = super(Doc, self).add_table(*(df.shape), style)
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                table.cell(i, j).text = str(df.iloc[i, j])
        self.beautify_table(table)
        return table

    def save(self, path_or_stream):
        super(Doc, self).save(path_or_stream)

    def move_table_after(self, table, paragraph):
        paragraph._p.addnext(table._tbl)

    def delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    def merge_cell(self, a: _Cell, b: _Cell):
        a.merge(b)
        if a.text.strip('\n') != b.text.strip('\n'):
            a.text = (a.text + b.text).replace('\n', '')

    def beautify_paragraph(self, p, size=0, bold=False, color=RGBColor(0, 0, 0), alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
        p.alignment = alignment
        if p.runs:
            run = p.runs[0]
        else:
            run = p.add_run(p.text)
        run.bold = bold
        run.font.size = Pt(size) if size else None
        run.font.color.rgb = color

    def beautify_table(self, table):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell in table.rows[0].cells:
            self.beautify_paragraph(cell.paragraphs[0], bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        border_map = {'w:sz': '12', 'w:val': 'single', 'w:color': '000000', 'w:space': '0', 'w:shadow': '12'}
        twidth = table._tblPr.first_child_found_in('w:tblW')
        twidth.set(ns.qn('w:w'), '6000')
        twidth.set(ns.qn('w:type'), 'pct')
        tborder = table._tblPr.first_child_found_in('w:tblBorders')
        if tborder is None:
            tborder = OxmlElement('w:tblBorders')
            table._tblPr.append(tborder)
        for edge in ['top', 'right', 'bottom', 'left']:  # 'insideH', 'insideV'
            element = tborder.find(ns.qn(f'w:{edge}'))
            if element is None:
                element = OxmlElement(f'w:{edge}')
                tborder.append(element)
            for k, v in border_map.items():
                element.set(ns.qn(k), v)

    def fill(self, table_map, paragraph_map, bold_names=None):
        bold_names = bold_names or []
        for p in self.paragraphs:
            if p.text.startswith('{{p.'):
                p.text = paragraph_map.get(p.text[4:-2], '')
            elif p.text.startswith('{{table.'):
                df = table_map.get(p.text[4:-2]).T.reset_index().T
                table = self.add_table(rows=0, cols=0, df=df)
                self.move_table_after(table, p)
                self.delete_paragraph(p)
        for t in self.tables:
            for r in t.rows:
                is_bold = r.cells[0].text in bold_names
                for c in r.cells:
                    if c.text.startswith('{{p.'):
                        c.text = paragraph_map.get(c.text[4:-2], c.text)
                    if is_bold:
                        self.beautify_paragraph(c.paragraphs[0], bold=True)
        print('over')
